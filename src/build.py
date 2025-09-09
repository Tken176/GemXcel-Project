#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Demo: Đọc các loại file document trong ứng dụng Python
Hỗ trợ: .docx, .doc, .txt, .md, .pdf, .xlsx, .xls
"""

import os
import sys
from pathlib import Path

def read_text_file(file_path):
    """Đọc file text thuần (.txt, .md, .py, v.v.)"""
    try:
        # Thử các encoding khác nhau
        encodings = ['utf-8', 'utf-8-sig', 'cp1252', 'latin1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                print(f"✅ Đọc file TEXT thành công với encoding: {encoding}")
                return content
            except UnicodeDecodeError:
                continue
        
        # Nếu tất cả encoding đều fail, dùng chardet
        try:
            import chardet
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding']
            
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
            print(f"✅ Đọc file TEXT với chardet encoding: {encoding}")
            return content
        except Exception as e:
            print(f"❌ Không thể đọc file text: {e}")
            return None
            
    except Exception as e:
        print(f"❌ Lỗi đọc file text: {e}")
        return None

def read_word_docx(file_path):
    """Đọc file Word .docx"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        content = []
        
        # Đọc từng paragraph
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # Đọc tables nếu có
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        result = "\n".join(content)
        print("✅ Đọc file DOCX thành công")
        return result
        
    except ImportError:
        print("❌ Cần cài đặt: pip install python-docx")
        return None
    except Exception as e:
        print(f"❌ Lỗi đọc file DOCX: {e}")
        return None

def read_word_with_mammoth(file_path):
    """Đọc file Word bằng mammoth (tốt hơn cho formatting)"""
    try:
        import mammoth
        
        with open(file_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value
            
            # Convert HTML to plain text (basic)
            import re
            text = re.sub('<[^<]+?>', '', html_content)
            text = text.replace('&nbsp;', ' ').replace('&amp;', '&')
            
        print("✅ Đọc file Word bằng mammoth thành công")
        return text
        
    except ImportError:
        print("❌ Cần cài đặt: pip install mammoth")
        return None
    except Exception as e:
        print(f"❌ Lỗi đọc file với mammoth: {e}")
        return None

def read_markdown_file(file_path):
    """Đọc file Markdown (.md)"""
    try:
        import markdown
        
        # Đọc file markdown
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
        
        # Convert markdown to HTML
        md = markdown.Markdown(extensions=['extra', 'codehilite'])
        html_content = md.convert(md_content)
        
        # Hoặc trả về raw markdown
        print("✅ Đọc file MARKDOWN thành công")
        return {
            'raw': md_content,
            'html': html_content
        }
        
    except ImportError:
        print("❌ Cần cài đặt: pip install markdown")
        return read_text_file(file_path)  # Fallback to text
    except Exception as e:
        print(f"❌ Lỗi đọc file markdown: {e}")
        return read_text_file(file_path)

def read_pdf_file(file_path):
    """Đọc file PDF"""
    try:
        import PyPDF2
        
        content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        content.append(f"--- Trang {page_num + 1} ---")
                        content.append(text)
                except Exception as e:
                    print(f"⚠️ Không đọc được trang {page_num + 1}: {e}")
        
        result = "\n".join(content)
        print("✅ Đọc file PDF thành công")
        return result
        
    except ImportError:
        print("❌ Cần cài đặt: pip install PyPDF2")
        return None
    except Exception as e:
        print(f"❌ Lỗi đọc file PDF: {e}")
        
        # Thử với pdfplumber
        try:
            import pdfplumber
            
            content = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        content.append(f"--- Trang {page_num + 1} ---")
                        content.append(text)
            
            result = "\n".join(content)
            print("✅ Đọc file PDF bằng pdfplumber thành công")
            return result
            
        except ImportError:
            print("❌ Cần cài đặt: pip install pdfplumber")
        except Exception as e2:
            print(f"❌ Lỗi đọc PDF với pdfplumber: {e2}")
        
        return None

def read_excel_file(file_path):
    """Đọc file Excel (.xlsx, .xls)"""
    try:
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.xlsx':
            import openpyxl
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content.append(f"=== Sheet: {sheet_name} ===")
                
                for row in sheet.iter_rows(values_only=True):
                    if any(cell for cell in row):  # Không bỏ qua row trống
                        row_text = []
                        for cell in row:
                            row_text.append(str(cell) if cell is not None else "")
                        content.append(" | ".join(row_text))
            
            result = "\n".join(content)
            print("✅ Đọc file XLSX thành công")
            return result
            
        elif file_ext == '.xls':
            import xlrd
            
            workbook = xlrd.open_workbook(file_path)
            content = []
            
            for sheet_name in workbook.sheet_names():
                sheet = workbook.sheet_by_name(sheet_name)
                content.append(f"=== Sheet: {sheet_name} ===")
                
                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    if any(str(cell).strip() for cell in row):
                        row_text = [str(cell) for cell in row]
                        content.append(" | ".join(row_text))
            
            result = "\n".join(content)
            print("✅ Đọc file XLS thành công")
            return result
            
    except ImportError as e:
        print(f"❌ Thiếu package: {e}")
        return None
    except Exception as e:
        print(f"❌ Lỗi đọc file Excel: {e}")
        return None

def detect_file_type(file_path):
    """Phát hiện loại file"""
    file_ext = Path(file_path).suffix.lower()
    
    file_types = {
        '.txt': 'text',
        '.md': 'markdown', 
        '.markdown': 'markdown',
        '.docx': 'word_docx',
        '.doc': 'word_doc',
        '.pdf': 'pdf',
        '.xlsx': 'excel',
        '.xls': 'excel',
        '.py': 'text',
        '.js': 'text',
        '.html': 'text',
        '.css': 'text',
        '.json': 'text',
        '.xml': 'text',
        '.csv': 'text'
    }
    
    return file_types.get(file_ext, 'unknown')

def read_document(file_path):
    """Hàm chính để đọc bất kỳ loại document nào"""
    if not os.path.exists(file_path):
        print(f"❌ File không tồn tại: {file_path}")
        return None
    
    file_type = detect_file_type(file_path)
    print(f"📄 Đang đọc file: {file_path}")
    print(f"🔍 Loại file: {file_type}")
    
    try:
        if file_type == 'text':
            return read_text_file(file_path)
        elif file_type == 'markdown':
            return read_markdown_file(file_path)
        elif file_type == 'word_docx':
            # Thử mammoth trước, fallback to python-docx
            result = read_word_with_mammoth(file_path)
            if result is None:
                result = read_word_docx(file_path)
            return result
        elif file_type == 'pdf':
            return read_pdf_file(file_path)
        elif file_type == 'excel':
            return read_excel_file(file_path)
        else:
            print(f"⚠️ Loại file không được hỗ trợ: {file_type}")
            # Thử đọc như text file
            return read_text_file(file_path)
            
    except Exception as e:
        print(f"❌ Lỗi đọc document: {e}")
        return None

# Demo usage
def demo():
    """Demo cách sử dụng"""
    print("=" * 60)
    print("📚 DOCUMENT READER DEMO")
    print("=" * 60)
    
    # Test files (tạo demo files nếu không có)
    test_files = [
        "demo.txt",
        "demo.md", 
        "demo.docx",
        "demo.pdf"
    ]
    
    for test_file in test_files:
        if os.path.exists(test_file):
            print(f"\n🔍 Testing: {test_file}")
            content = read_document(test_file)
            if content:
                # Hiển thị 200 ký tự đầu
                preview = content[:200] + "..." if len(content) > 200 else content
                print(f"📖 Nội dung preview:\n{preview}")
            else:
                print("❌ Không đọc được file")
        else:
            print(f"⚠️ File {test_file} không tồn tại, bỏ qua...")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        # Đọc file từ command line argument
        file_path = sys.argv[1]
        content = read_document(file_path)
        if content:
            print("\n" + "="*60)
            print("📖 NỘI DUNG FILE:")
            print("="*60)
            print(content)
    else:
        # Chạy demo
        demo()
        print("\n💡 Cách sử dụng: python document_reader.py <đường_dẫn_file>")
