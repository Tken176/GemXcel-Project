import os
import json
import re
import time
import hashlib
import signal
import requests
import chardet
from typing import List, Dict, Optional, Tuple, Set 
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
import config
import random


class TimeoutException(Exception):
    """Custom timeout exception"""
    pass


def timeout_handler(signum, frame):
    """Signal handler for timeout"""
    raise TimeoutException("Operation timed out")


class GeminiClient:
    """Gemini AI Client để thay thế g4f"""
    
    def __init__(self):
        self.api_key = "INPUT_YOUR_API"  
        self.base_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
        self.session = requests.Session()
        
    def chat_completions_create(self, messages, temperature=0.7, **kwargs):
        """Tương thích với interface của g4f client"""
        # Lấy nội dung từ messages (format của OpenAI/g4f)
        if isinstance(messages, list) and len(messages) > 0:
            prompt = messages[-1].get("content", "")
        else:
            prompt = str(messages)
        
        headers = {
            "Content-Type": "application/json",
            "X-goog-api-key": self.api_key
        }
        
        data = {
            "contents": [
                {
                    "parts": [
                        {"text": prompt}
                    ]
                }
            ],
            "generationConfig": {
                "temperature": temperature,
                "maxOutputTokens": 4096,
                "topP": 0.8,
                "topK": 10
            }
        }
        
        try:
            response = self.session.post(
                self.base_url, 
                headers=headers, 
                json=data, 
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                try:
                    content = result['candidates'][0]['content']['parts'][0]['text']
                    # Trả về object tương thích với g4f response format
                    return GeminiResponse(content)
                except (KeyError, IndexError):
                    raise Exception("Không nhận được phản hồi từ Gemini API")
            else:
                raise Exception(f"Gemini API error {response.status_code}: {response.text}")
                
        except requests.exceptions.Timeout:
            raise Exception("Gemini API timeout")
        except requests.exceptions.RequestException as e:
            raise Exception(f"Gemini API request error: {e}")


class GeminiResponse:
    """Response object tương thích với g4f format"""
    
    def __init__(self, content):
        self.choices = [GeminiChoice(content)]


class GeminiChoice:
    """Choice object tương thích với g4f format"""
    
    def __init__(self, content):
        self.message = GeminiMessage(content)


class GeminiMessage:
    """Message object tương thích với g4f format"""
    
    def __init__(self, content):
        self.content = content


class GeminiChatCompletions:
    """Chat completions wrapper tương thích với g4f"""
    
    def __init__(self, client):
        self.client = client
        
    def create(self, model=None, messages=None, temperature=0.7, **kwargs):
        return self.client.chat_completions_create(
            messages=messages, 
            temperature=temperature, 
            **kwargs
        )


class GeminiChat:
    """Chat wrapper tương thích với g4f"""
    
    def __init__(self, client):
        self.completions = GeminiChatCompletions(client)


class ContentProcessor:
    def __init__(self, lessons_count: int = 5, questions_per_lesson: int = 6, quiz_questions: int = 10):
        """
        Khởi tạo ContentProcessor với cấu hình linh hoạt

        Args:
            lessons_count: Số lượng bài học (mặc định 5, có thể tăng lên 20)
            questions_per_lesson: Số câu hỏi mỗi bài (mặc định 6, có thể tăng lên 15)
            quiz_questions: Số câu hỏi quiz tổng hợp (mặc định 10, có thể tăng lên 50)
        """
        # Khởi tạo Gemini client thay vì g4f
        self.gemini_client = GeminiClient()
        self.client = type('Client', (), {
            'chat': GeminiChat(self.gemini_client)
        })()
        
        self.lessons_path = config.LESSON_DATA_FILE_PATH
        self.quiz_path = config.QUIZ_DATA_FILE_PATH

        self.supported_formats = ['.txt', '.docx', '.md', '.pdf']
        self.max_retries = 3
        self.timeout = 45  # Giảm timeout xuống 45s cho mỗi request
        self.total_timeout = 300  # Timeout tổng cho toàn bộ quá trình: 5 phút

        # Cấu hình linh hoạt
        self.lessons_count = max(1, min(lessons_count, 20))
        self.questions_per_lesson = max(3, min(questions_per_lesson, 15))
        self.quiz_questions = max(5, min(quiz_questions, 50))

        # Cache đơn giản
        self.content_cache = {}

        # Đảm bảo thư mục tồn tại
        os.makedirs(config.ASSETS_DIR, exist_ok=True)

        # Flags for graceful shutdown
        self.should_stop = False

        # Tracking để tránh câu hỏi lặp
        self.used_questions = set()
        self.used_quiz_questions = set()

    # =========================
    # Helpers
    # =========================

    def is_supported(self, file_path: str) -> bool:
        """Kiểm tra xem file có được hỗ trợ không"""
        return any(file_path.lower().endswith(ext) for ext in self.supported_formats)

    def process_file(self, file_path: str) -> Tuple[bool, str]:
        """Xử lý file chính với timeout và error handling toàn diện"""
        start_time = time.time()
        
        # Reset tracking sets cho mỗi file mới
        self.used_questions.clear()
        self.used_quiz_questions.clear()
        
        # Set up timeout cho toàn bộ quá trình
        if hasattr(signal, 'SIGALRM'):  # Unix systems
            signal.signal(signal.SIGALRM, timeout_handler)
            signal.alarm(self.total_timeout)

        try:
            # Kiểm tra định dạng file
            if not self.is_supported(file_path):
                return False, f"Định dạng file không được hỗ trợ. Các định dạng hỗ trợ: {', '.join(self.supported_formats)}"

            # Kiểm tra file tồn tại
            if not os.path.exists(file_path):
                return False, f"File không tồn tại: {file_path}"

            # Đọc và validate file
            try:
                content = self.read_file_content(file_path)
            except Exception as e:
                return False, f"Không thể đọc file: {e}"
            if not content or not content.strip():
                return False, "Không thể đọc nội dung từ file (file rỗng hoặc định dạng không đúng)."

            if len(content.strip()) < 200:
                return False, "Nội dung file quá ngắn để xử lý (cần ít nhất 200 ký tự)."

            # Tiền xử lý nội dung
            content = self._preprocess_content(content)

            # Kiểm tra cache
            content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
            if content_hash in self.content_cache:
                lessons, quiz_questions = self.content_cache[content_hash]
            else:
                # Chia nội dung thành các phần
                chunks = self._intelligent_chunking(content, target_chunks=self.lessons_count)

                # Tạo bài học với timeout
                lessons = self.generate_lessons_with_timeout(chunks)
                
                if not lessons:
                    return False, "Không thể tạo bài học từ nội dung này."

                # Lọc None và validate
                lessons = [lesson for lesson in lessons if lesson is not None]
                if not lessons:
                    return False, "Không thể tạo bài học hợp lệ từ nội dung này."

                # Tạo quiz với timeout
                quiz_questions = self.generate_quiz_with_timeout(lessons, total_questions=self.quiz_questions)

                # Cache kết quả
                self.content_cache[content_hash] = (lessons, quiz_questions)

            # Lưu file với error handling
            success = self._safe_save_data_files(lessons, quiz_questions)
            if not success:
                return False, "Không thể lưu dữ liệu. Vui lòng kiểm tra quyền ghi file."

            elapsed_time = time.time() - start_time
            total_questions = sum(len(lesson.get('questions', [])) for lesson in lessons)
            total_quiz = sum(len(q) for q in quiz_questions.values()) if isinstance(quiz_questions, dict) else 0

            return True, f"Xử lý thành công trong {elapsed_time:.1f}s! Đã tạo {len(lessons)} bài học ({total_questions} câu hỏi) và {total_quiz} câu hỏi quiz không trùng lặp."

        except TimeoutException:
            elapsed_time = time.time() - start_time
            return False, f"Timeout sau {elapsed_time:.1f}s! Quá trình xử lý mất quá nhiều thời gian. Vui lòng thử lại với nội dung ngắn hơn."
        
        except Exception as e:
            elapsed_time = time.time() - start_time
            error_msg = str(e)
            if "timeout" in error_msg.lower():
                return False, f"Timeout sau {elapsed_time:.1f}s! Server AI phản hồi chậm. Vui lòng thử lại."
            else:
                return False, f"Lỗi sau {elapsed_time:.1f}s: {error_msg}"
        
        finally:
            # Cancel timeout alarm
            if hasattr(signal, 'SIGALRM'):
                signal.alarm(0)

    # =========================
    # Đọc & Tiền xử lý nội dung với encoding detection
    # =========================

    def _detect_encoding(self, file_path: str) -> str:
        """Phát hiện encoding của file"""
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # Đọc 10KB đầu để detect
                result = chardet.detect(raw_data)
                confidence = result.get('confidence', 0)
                encoding = result.get('encoding', 'utf-8')
                
                # Nếu confidence thấp hoặc encoding không xác định, thử các encoding phổ biến
                if confidence < 0.7 or not encoding:
                    return 'utf-8'
                    
                # Mapping một số encoding thường gặp lỗi
                encoding_map = {
                    'utf-16': 'utf-16',
                    'utf-16le': 'utf-16-le',
                    'utf-16be': 'utf-16-be', 
                    'windows-1252': 'cp1252',
                    'iso-8859-1': 'latin-1'
                }
                
                return encoding_map.get(encoding.lower(), encoding)
        except Exception:
            return 'utf-8'

    def _preprocess_content(self, content: str) -> str:
        """Tiền xử lý nội dung nhưng giữ ký tự đặc biệt cần thiết"""
        # Gom khoảng trắng: giữ lại xuống dòng 2 dấu cách tối thiểu giữa đoạn
        content = re.sub(r'\n\s*\n', '\n\n', content)
        content = re.sub(r'[ \t]+', ' ', content)

        # Loại bỏ control characters nhưng giữ \n, \t
        cleaned_chars = []
        for ch in content:
            if ch == '\n' or ch == '\t' or ch.isprintable():
                cleaned_chars.append(ch)
        content = ''.join(cleaned_chars)

        return content.strip()

    def read_file_content(self, file_path: str) -> Optional[str]:
        """Đọc nội dung từ file với xử lý lỗi tốt hơn và encoding detection"""
        try:
            if not self.is_supported(file_path):
                raise ValueError(f"Định dạng file không được hỗ trợ: {file_path}")

            if file_path.lower().endswith('.txt'):
                # Tự động detect encoding
                detected_encoding = self._detect_encoding(file_path)
                
                # Danh sách encoding fallback
                encodings = [detected_encoding, 'utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 
                           'cp1252', 'latin-1', 'gbk', 'big5']
                
                # Thử từng encoding
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                            if content.strip():  # Đảm bảo có nội dung
                                return content
                    except (UnicodeDecodeError, UnicodeError):
                        continue
                
                # Nếu tất cả encoding đều fail, thử đọc binary và decode từng phần
                try:
                    with open(file_path, 'rb') as f:
                        raw_data = f.read()
                    
                    # Thử decode với errors='ignore' hoặc 'replace'
                    for encoding in ['utf-8', 'cp1252', 'latin-1']:
                        try:
                            content = raw_data.decode(encoding, errors='replace')
                            if content.strip() and not self._is_binary_garbage(content):
                                return content
                        except Exception:
                            continue
                            
                except Exception:
                    pass
                
                raise ValueError("Không thể đọc file .txt với các encoding thông dụng")

            elif file_path.lower().endswith('.docx'):
                try:
                    import docx
                    
                    # Kiểm tra file có phải DOCX hợp lệ không
                    if not os.path.exists(file_path):
                        raise ValueError("File không tồn tại")
                    
                    # Kiểm tra file size
                    if os.path.getsize(file_path) == 0:
                        raise ValueError("File DOCX rỗng")
                    
                    # Thử đọc document
                    try:
                        doc = docx.Document(file_path)
                    except Exception as e:
                        raise ValueError(f"File DOCX bị lỗi hoặc không hợp lệ: {e}")
                    
                    # Trích xuất text từ paragraphs
                    paragraphs_text = []
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:  # Chỉ lấy paragraph có nội dung
                            paragraphs_text.append(text)
                    
                    # Trích xuất text từ tables (nếu có)
                    tables_text = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text:
                                    row_text.append(cell_text)
                            if row_text:
                                tables_text.append(" | ".join(row_text))
                    
                    # Kết hợp tất cả text
                    all_text = paragraphs_text + tables_text
                    content = "\n\n".join(all_text)
                    
                    if not content.strip():
                        raise ValueError("File DOCX không chứa text có thể đọc được")
                    
                    # Kiểm tra content có ý nghĩa không (không phải binary garbage)
                    if self._is_binary_garbage(content):
                        raise ValueError("Nội dung DOCX chứa dữ liệu binary không hợp lệ")
                    
                    return content
                    
                except ImportError:
                    raise ValueError("Cần cài đặt python-docx để đọc file .docx: pip install python-docx")
                except Exception as e:
                    # Fallback: thử đọc như ZIP để debug
                    try:
                        import zipfile
                        if zipfile.is_zipfile(file_path):
                            raise ValueError(f"File DOCX hợp lệ nhưng không đọc được nội dung: {e}")
                        else:
                            raise ValueError(f"File không phải DOCX hợp lệ: {e}")
                    except ImportError:
                        raise ValueError(f"Lỗi đọc file DOCX: {e}")

            elif file_path.lower().endswith('.md'):
                # Detect encoding cho file Markdown
                detected_encoding = self._detect_encoding(file_path)
                encodings = [detected_encoding, 'utf-8', 'cp1252', 'latin-1']
                
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                            if content.strip():
                                # Loại bỏ markdown syntax cơ bản nhưng giữ text
                                content = re.sub(r'(^|\n)#{1,6}\s*', r'\1', content)  # Headers
                                content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)    # Bold
                                content = re.sub(r'\*(.*?)\*', r'\1', content)        # Italic
                                content = re.sub(r'`{1,3}([^`]+)`{1,3}', r'\1', content)  # Inline code
                                return content
                    except (UnicodeDecodeError, UnicodeError):
                        continue
                        
                raise ValueError("File Markdown không đọc được với các encoding thông dụng")

            elif file_path.lower().endswith('.pdf'):
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        pages_text = []
                        for page in reader.pages:
                            try:
                                text = page.extract_text()
                                if text and text.strip():
                                    pages_text.append(text)
                            except Exception:
                                continue
                        content = "\n".join(pages_text)
                        if not content.strip():
                            raise ValueError("Không thể trích xuất text từ PDF")
                        return content
                except ImportError:
                    raise ValueError("Cần cài đặt PyPDF2 để đọc file .pdf: pip install PyPDF2")

        except Exception as e:
            print(f"Lỗi đọc file {file_path}: {e}")
            return None

    def _is_binary_garbage(self, content: str) -> bool:
        """Kiểm tra xem content có phải binary garbage không"""
        if not content:
            return True
        
        # Đếm ký tự không thể in được
        printable_chars = sum(1 for c in content if c.isprintable() or c.isspace())
        total_chars = len(content)
        
        # Nếu ít hơn 80% là ký tự có thể in được thì có thể là binary
        if total_chars > 0:
            printable_ratio = printable_chars / total_chars
            return printable_ratio < 0.8
        
        return True

    # =========================
    # Chunking
    # =========================

    def _intelligent_chunking(self, content: str, target_chunks: int = 5) -> List[str]:
        """Chia content thành các phần thông minh hơn, ưu tiên theo đoạn văn"""
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]

        if len(paragraphs) <= target_chunks:
            # Nếu ít đoạn văn, chia theo từ
            words = content.split()
            chunk_size = max(80, len(words) // max(1, target_chunks))  # Tối thiểu 80 từ/chunk
            chunks = []
            for i in range(target_chunks):
                start_idx = i * chunk_size
                end_idx = min((i + 1) * chunk_size, len(words))
                if start_idx < len(words):
                    chunk = ' '.join(words[start_idx:end_idx])
                    if chunk.strip():
                        chunks.append(chunk)
            return chunks
        else:
            # Gom số đoạn văn về gần target_chunks
            base = len(paragraphs) // target_chunks
            rem = len(paragraphs) % target_chunks
            chunks = []
            start = 0
            for i in range(target_chunks):
                take = base + (1 if i < rem else 0)
                end = start + take
                chunk = '\n\n'.join(paragraphs[start:end]).strip()
                if chunk:
                    chunks.append(chunk)
                start = end
            return chunks

    # =========================
    # Sinh bài học với timeout và tránh lặp
    # =========================

    def generate_lessons_with_timeout(self, chunks: List[str]) -> List[Dict]:
        """Tạo bài học với timeout và error handling toàn diện"""
        results = [None] * len(chunks)
        completed_count = 0

        def _lesson_callback(future, lesson_num):
            nonlocal completed_count
            try:
                result = future.result(timeout=self.timeout)
                results[lesson_num - 1] = result
                completed_count += 1
            except Exception as e:
                results[lesson_num - 1] = self._create_fallback_lesson(
                    lesson_num,
                    chunks[lesson_num - 1] if lesson_num - 1 < len(chunks) else ""
                )
                completed_count += 1

        with ThreadPoolExecutor(max_workers=min(3, len(chunks))) as executor:
            futures = []
            for i, chunk in enumerate(chunks):
                if self.should_stop:
                    break
                future = executor.submit(self._generate_single_lesson_safe, chunk, i + 1)
                future.add_done_callback(lambda f, num=i + 1: _lesson_callback(f, num))
                futures.append(future)

            # Chờ với timeout
            try:
                for _ in as_completed(futures, timeout=self.timeout * len(chunks)):
                    if self.should_stop:
                        break
            except TimeoutError:
                pass

        return [r for r in results if r is not None]

    def _generate_unique_question_hash(self, question: str, choices: List[str]) -> str:
        """Tạo hash để identify câu hỏi duy nhất"""
        # Chuẩn hóa câu hỏi và đáp án
        normalized_question = re.sub(r'\s+', ' ', question.strip().lower())
        normalized_choices = [re.sub(r'\s+', ' ', choice.strip().lower()) for choice in choices]
        
        # Tạo hash từ câu hỏi và đáp án
        combined = normalized_question + '|' + '|'.join(normalized_choices)
        return hashlib.md5(combined.encode('utf-8')).hexdigest()

    def _is_question_unique(self, question: str, choices: List[str], used_set: Set[str]) -> bool:
        """Kiểm tra câu hỏi có trùng lặp không"""
        question_hash = self._generate_unique_question_hash(question, choices)
        return question_hash not in used_set

    def _add_question_to_used(self, question: str, choices: List[str], used_set: Set[str]):
        """Thêm câu hỏi vào danh sách đã dùng"""
        question_hash = self._generate_unique_question_hash(question, choices)
        used_set.add(question_hash)

    def _generate_single_lesson_safe(self, chunk: str, lesson_number: int) -> Dict:
        """Tạo một bài học với timeout và error handling an toàn"""
        for attempt in range(self.max_retries):
            try:
                if self.should_stop:
                    break

                # Giới hạn độ dài chunk để tránh vượt quá token limit
                max_chunk_length = 1200  # Giảm xuống để tăng tốc độ
                chunk_for_prompt = (chunk[:max_chunk_length] + "...") if len(chunk) > max_chunk_length else chunk

                prompt = f"""
Tạo bài học số {lesson_number} từ nội dung sau:
{chunk_for_prompt}

Yêu cầu:
0. Không bị lỗi charmap
1. Tiêu đề ngắn gọn, phản ánh nội dung chính
2. Nội dung tóm tắt 230-350 từ, tập trung vào ý quan trọng
3. {self.questions_per_lesson} câu hỏi với độ khó phân bố đều:
   - {self.questions_per_lesson // 3} câu dễ (easy)
   - {self.questions_per_lesson // 3} câu trung bình (medium)
   - {self.questions_per_lesson - 2 * (self.questions_per_lesson // 3)} câu khó (hard)
4. Mỗi câu hỏi có:
   - Câu hỏi rõ ràng, liên quan trực tiếp đến nội dung
   - 4 lựa chọn A, B, C, D (mỗi đáp án tối đa 50 ký tự)
   - Đáp án đúng là chỉ số 0-3 (0 = A, 1 = B, 2 = C, 3 = D)
   - KHÔNG tạo câu hỏi mẫu hay placeholder
   - Câu hỏi phải ĐA DẠNG và KHÔNG TRÙNG LẶP
   - Tránh các câu hỏi chung chung như "Nội dung chính là gì?"

Trả về **chỉ JSON** hợp lệ theo mẫu:
{{
  "name": "Bài {lesson_number}",
  "title": "[Tiêu đề cụ thể về nội dung]",
  "content": "[Nội dung tóm tắt chi tiết]",
  "questions": [
    {{
      "question": "[Câu hỏi chi tiết về nội dung]",
      "choices": ["A. ...", "B. ...", "C. ...", "D. ..."],
      "correct_answer": 0,
      "difficulty": "easy"
    }}
  ]
}}
"""

                # Gọi Gemini API với timeout
                response = self.client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.6,  # Giảm temperature để tăng độ ổn định
                )

                response_text = (response.choices[0].message.content or "").strip()

                if not response_text:
                    raise ValueError("Response rỗng từ AI")

                # Parse JSON an toàn
                lesson = self._safe_parse_json(response_text)
                if not lesson:
                    raise ValueError("Không parse được JSON từ response")

                return self._validate_lesson(lesson, lesson_number, chunk)

            except Exception as e:
                if attempt == self.max_retries - 1:
                    return self._create_fallback_lesson(lesson_number, chunk)
                time.sleep(min(2 ** attempt, 5))  # Exponential backoff với max 5s

    def _safe_parse_json(self, text: str) -> Optional[Dict]:
        # Thử parse trực tiếp
        try:
            parsed = json.loads(text)
            # Nếu là list và chứa dict(s), trả phần tử đầu tiên nếu structure hợp lệ
            if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
                if self._is_valid_lesson_structure(parsed[0]):
                    return parsed[0]
                # else maybe list of lessons but different fields: try first that matches
                for item in parsed:
                    if isinstance(item, dict) and self._is_valid_lesson_structure(item):
                        return item
                # fallback: return first dict
                return parsed[0]
            if isinstance(parsed, dict) and self._is_valid_lesson_structure(parsed):
                return parsed
        except json.JSONDecodeError:
            pass

        # Tìm JSON object trong text
        json_candidates = self._extract_json_objects(text)
        for candidate in json_candidates:
            try:
                parsed = json.loads(candidate)
                if isinstance(parsed, dict) and self._is_valid_lesson_structure(parsed):
                    return parsed
                if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
                    return parsed[0]
            except json.JSONDecodeError:
                continue

        # Thử cleanup code fences + strip text trước/sau
        try:
            cleaned = re.sub(r'```(?:json)?\s*', '', text)
            cleaned = re.sub(r'```(?:json)?\s*', '', cleaned).strip(" \n`")
            parsed = json.loads(cleaned)
            if isinstance(parsed, dict) and self._is_valid_lesson_structure(parsed):
                return parsed
            if isinstance(parsed, list) and parsed:
                return parsed[0]
        except Exception:
            pass

        return None

    def _is_valid_lesson_structure(self, data: Dict) -> bool:
        """Kiểm tra cấu trúc lesson có hợp lệ không"""
        required_fields = ['name', 'title', 'content', 'questions']
        if not all(field in data for field in required_fields):
            return False
        
        if not isinstance(data['questions'], list):
            return False
            
        return True

    def _extract_json_objects(self, text: str) -> List[str]:
        """
        Quét chuỗi để trích xuất các JSON object cấp cao nhất bằng cách đếm dấu ngoặc.
        """
        results = []
        stack = 0
        in_string = False
        escape = False
        start_idx = None

        for i, ch in enumerate(text):
            if in_string:
                if escape:
                    escape = False
                elif ch == '\\':
                    escape = True
                elif ch == '"':
                    in_string = False
                continue
            else:
                if ch == '"':
                    in_string = True
                    continue
                if ch == '{':
                    if stack == 0:
                        start_idx = i
                    stack += 1
                elif ch == '}':
                    if stack > 0:
                        stack -= 1
                        if stack == 0 and start_idx is not None:
                            results.append(text[start_idx:i + 1])
                            start_idx = None

        return results

    # =========================
    # Validate & Fallback với tránh lặp
    # =========================

    def _validate_lesson(self, lesson: Dict, lesson_number: int, chunk: str) -> Dict:
        """Đảm bảo bài học có đủ câu hỏi và dữ liệu hợp lệ, tránh trùng lặp"""
        # Trường bắt buộc
        lesson.setdefault("name", f"Bài {lesson_number}")
        lesson.setdefault("title", f"Bài học {lesson_number}")
        lesson.setdefault("content", chunk[:300] + "..." if len(chunk) > 300 else chunk)

        # Câu hỏi
        if "questions" not in lesson or not isinstance(lesson["questions"], list):
            lesson["questions"] = []

        # Lọc câu hỏi trùng lặp và validate
        unique_questions = []
        for q in lesson["questions"]:
            if not isinstance(q, dict):
                continue
                
            question_text = q.get("question", "")
            choices = q.get("choices", [])
            
            if isinstance(choices, list) and len(choices) == 4:
                # Kiểm tra unique
                if self._is_question_unique(question_text, choices, self.used_questions):
                    self._add_question_to_used(question_text, choices, self.used_questions)
                    unique_questions.append(q)

        lesson["questions"] = unique_questions

        # Đảm bảo đủ số câu hỏi
        while len(lesson["questions"]) < self.questions_per_lesson:
            new_question = self._create_default_question(
                len(lesson["questions"]) + 1,
                lesson_number,
                chunk
            )
            
            # Kiểm tra unique cho câu hỏi mới tạo
            if self._is_question_unique(new_question["question"], new_question["choices"], self.used_questions):
                self._add_question_to_used(new_question["question"], new_question["choices"], self.used_questions)
                lesson["questions"].append(new_question)
            else:
                # Tạo câu hỏi khác nếu bị trùng
                new_question = self._create_varied_question(
                    len(lesson["questions"]) + 1,
                    lesson_number,
                    chunk
                )
                self._add_question_to_used(new_question["question"], new_question["choices"], self.used_questions)
                lesson["questions"].append(new_question)

        # Giới hạn số câu hỏi
        lesson["questions"] = lesson["questions"][:self.questions_per_lesson]

        # Validate từng câu hỏi
        for i, q in enumerate(lesson["questions"]):
            if not isinstance(q, dict):
                lesson["questions"][i] = self._create_default_question(i + 1, lesson_number, chunk)
                continue

            q.setdefault("question", f"Câu hỏi {i + 1} bài {lesson_number}")
            q.setdefault("choices", ["A. Lựa chọn A", "B. Lựa chọn B", "C. Lựa chọn C", "D. Lựa chọn D"])
            q.setdefault("correct_answer", 0)
            q.setdefault("difficulty", "medium")

            # Đảm bảo có đủ 4 lựa chọn
            if not isinstance(q["choices"], list) or len(q["choices"]) != 4:
                q["choices"] = ["A. Lựa chọn A", "B. Lựa chọn B", "C. Lựa chọn C", "D. Lựa chọn D"]
            
            # Giới hạn độ dài đáp án 50 ký tự
            for idx, choice in enumerate(q["choices"]):
                if len(choice) > 50:
                    prefix = choice[:2] if choice.startswith(('A.', 'B.', 'C.', 'D.')) else ""
                    content = choice[2:].strip() if prefix else choice
                    q["choices"][idx] = prefix + " " + content[:47-len(prefix)] if prefix else content[:50]

            # Đảm bảo correct_answer hợp lệ
            if not isinstance(q["correct_answer"], int) or q["correct_answer"] < 0 or q["correct_answer"] > 3:
                q["correct_answer"] = 0

            # Chuẩn hóa difficulty
            if q.get("difficulty") not in {"easy", "medium", "hard"}:
                q["difficulty"] = "medium"

        return lesson

    def _create_varied_question(self, q_num: int, lesson_num: int, content: str) -> Dict:
        """Tạo câu hỏi đa dạng để tránh trùng lặp"""
        difficulties = ["easy", "medium", "hard"]
        difficulty = difficulties[(q_num - 1) % 3]

        # Lấy mẫu nội dung để ghép vào câu hỏi
        words = content.split()
        sample_start = " ".join(words[:10]) if len(words) >= 10 else " ".join(words)
        sample_mid = " ".join(words[len(words)//2:len(words)//2+10]) if len(words) >= 20 else sample_start
        
        correct_ans = random.randint(0, 3)
        
        # Tạo câu hỏi đa dạng hơn
        question_templates = [
            f"Theo bài {lesson_num}, yếu tố quan trọng được nhấn mạnh là gì?",
            f"Bài {lesson_num} phân tích vấn đề nào sau đây?",
            f"Khái niệm chủ đạo trong bài {lesson_num} liên quan đến?",
            f"Nội dung bài {lesson_num} tập trung giải thích về?",
            f"Điểm cốt lõi của bài {lesson_num} là gì?",
            f"Theo nội dung bài {lesson_num}, đặc điểm chính là?",
            f"Bài {lesson_num} đề cập đến khía cạnh nào?",
            f"Thông tin then chốt trong bài {lesson_num} là?"
        ]
        
        # Tạo đáp án đa dạng
        answer_templates = [
            ["Nguyên lý cơ bản", "Phương pháp ứng dụng", "Quy trình thực hiện", "Kết quả đạt được"],
            ["Lý thuyết nền tảng", "Thực tiễn áp dụng", "Kinh nghiệm rút ra", "Hướng phát triển"],
            ["Khái niệm chính", "Đặc điểm nổi bật", "Ý nghĩa quan trọng", "Tác động tích cực"],
            ["Nội dung cốt lõi", "Yếu tố then chốt", "Giá trị thực tiễn", "Bài học kinh nghiệm"],
            ["Điểm mấu chốt", "Khía cạnh quan trọng", "Vấn đề cần lưu ý", "Hướng tiếp cận mới"],
        ]
        
        question = random.choice(question_templates)
        answer_set = random.choice(answer_templates)
        
        choices = []
        for i, answer in enumerate(answer_set):
            choices.append(f"{'ABCD'[i]}. {answer}")
        
        return {
            "question": question,
            "choices": choices,
            "correct_answer": correct_ans,
            "difficulty": difficulty
        }

    def _create_default_question(self, q_num: int, lesson_num: int, content: str) -> Dict:
        """Tạo câu hỏi mặc định từ nội dung"""
        difficulties = ["easy", "medium", "hard"]
        difficulty = difficulties[(q_num - 1) % 3]

        # Lấy mẫu nội dung để ghép vào câu hỏi
        words = content.split()
        sample = " ".join(words[:15]) if words else ""

        correct_ans = random.randint(0, 3)
        question_options = [
            f"Nội dung bài {lesson_num} chủ yếu đề cập đến vấn đề gì?",
            f"Ý chính của bài {lesson_num} là điều gì?",
            f"Bài {lesson_num} tập trung giải thích về?",
            f"Thông tin quan trọng nhất trong bài {lesson_num} là?"
        ]
        
        return {
            "question": random.choice(question_options),
            "choices": [
                f"A. {'Nội dung chính' if correct_ans == 0 else 'Thông tin bổ sung'}",
                f"B. {'Ý tưởng chính' if correct_ans == 1 else 'Ví dụ minh họa'}",
                f"C. {'Khái niệm cốt lõi' if correct_ans == 2 else 'Chi tiết kỹ thuật'}",
                f"D. {'Điểm quan trọng' if correct_ans == 3 else 'Kết luận chung'}"
            ],
            "correct_answer": correct_ans,
            "difficulty": difficulty
        }

    def _create_fallback_lesson(self, lesson_number: int, chunk: str) -> Dict:
        """Tạo bài học dự phòng chất lượng ổn"""
        content = chunk[:350] + "..." if len(chunk) > 350 else chunk

        # Tạo tiêu đề từ 8 từ đầu tiên
        words = chunk.split()[:8]
        title = " ".join(words) if words else f"Bài học {lesson_number}"
        if len(title) > 45:
            title = title[:42] + "..."

        questions = []
        for i in range(self.questions_per_lesson):
            new_question = self._create_varied_question(i + 1, lesson_number, chunk)
            # Đảm bảo unique
            if self._is_question_unique(new_question["question"], new_question["choices"], self.used_questions):
                self._add_question_to_used(new_question["question"], new_question["choices"], self.used_questions)
                questions.append(new_question)

        # Nếu vẫn thiếu câu hỏi, tạo thêm với template khác
        while len(questions) < self.questions_per_lesson:
            fallback_question = self._create_default_question(len(questions) + 1, lesson_number, chunk)
            questions.append(fallback_question)

        return {
            "name": f"Bài {lesson_number}",
            "title": title,
            "content": content,
            "questions": questions
        }

    # =========================
    # Quiz tổng hợp với timeout và tránh lặp
    # =========================

    def generate_quiz_with_timeout(self, lessons: List[Dict], total_questions: int = 10) -> Dict:
        """Tạo quiz tổng hợp với timeout"""
        try:
            return self.generate_quiz(lessons, total_questions)
        except Exception as e:
            return self._create_fallback_quiz(lessons, total_questions)

    def generate_quiz(self, lessons: List[Dict], total_questions: int = 10) -> Dict:
        """Tạo quiz tổng hợp với phân bố độ khó hợp lý (40/40/20) và tránh trùng lặp"""
        quiz = {"easy": [], "medium": [], "hard": []}

        # Phân bố câu hỏi theo tỷ lệ: 40% dễ, 40% trung bình, 20% khó
        easy_count = int(total_questions * 0.4)
        medium_count = int(total_questions * 0.4)
        hard_count = total_questions - easy_count - medium_count

        question_id = 1

        # Tạo câu hỏi dễ
        for _ in range(easy_count):
            question = self._generate_unique_quiz_question(question_id, "easy", lessons)
            if question:
                quiz["easy"].append(question)
                question_id += 1

        # Tạo câu hỏi trung bình
        for _ in range(medium_count):
            question = self._generate_unique_quiz_question(question_id, "medium", lessons)
            if question:
                quiz["medium"].append(question)
                question_id += 1

        # Tạo câu hỏi khó
        for _ in range(hard_count):
            question = self._generate_unique_quiz_question(question_id, "hard", lessons)
            if question:
                quiz["hard"].append(question)
                question_id += 1

        return quiz

    def _generate_unique_quiz_question(self, q_id: int, difficulty: str, lessons: List[Dict]) -> Dict:
        """Tạo câu hỏi quiz unique từ các bài học"""
        max_attempts = 10
        for attempt in range(max_attempts):
            question = self._generate_quiz_question(q_id, difficulty, lessons)
            
            # Kiểm tra unique
            if self._is_question_unique(question["question"], question["choices"], self.used_quiz_questions):
                self._add_question_to_used(question["question"], question["choices"], self.used_quiz_questions)
                return question
            
            # Thử tạo câu hỏi khác nếu bị trùng
            question = self._generate_varied_quiz_question(q_id, difficulty, lessons, attempt)
            if self._is_question_unique(question["question"], question["choices"], self.used_quiz_questions):
                self._add_question_to_used(question["question"], question["choices"], self.used_quiz_questions)
                return question
        
        # Fallback nếu không tạo được unique question
        return self._generate_quiz_question(q_id, difficulty, lessons)

    def _generate_varied_quiz_question(self, q_id: int, difficulty: str, lessons: List[Dict], variation: int) -> Dict:
        """Tạo câu hỏi quiz đa dạng để tránh trùng lặp"""
        import random
        lesson = random.choice(lessons) if lessons else None

        if not lesson:
            return self._generate_default_quiz_question(q_id, difficulty)

        content = lesson.get("content", "")
        title = lesson.get("title", f"Bài {q_id}")
        
        # Các template câu hỏi đa dạng hơn
        question_templates = [
            f"Theo {title}, khái niệm cơ bản được giới thiệu là gì?",
            f"Điểm đặc trưng của nội dung trong {title} là?",
            f"{title} phân tích vấn đề theo hướng nào?",
            f"Yếu tố được nhấn mạnh trong {title} là gì?",
            f"Nội dung {title} tập trung vào khía cạnh nào?",
            f"Theo {title}, nguyên tắc quan trọng là?",
            f"{title} đề xuất phương pháp gì?",
            f"Kết luận chính của {title} là gì?",
            f"Ý tưởng trung tâm trong {title} liên quan đến?",
            f"{title} giải thích hiện tượng như thế nào?"
        ]
        
        # Tạo đáp án đa dạng từ nội dung
        words = content.split()
        correct_answer = (variation + q_id) % 4  # Đảm bảo đáp án đúng đa dạng
        
        answer_pools = [
            ["Khái niệm nền tảng", "Phương pháp cơ bản", "Nguyên lý vận hành", "Quy trình thực hiện"],
            ["Đặc điểm chính yếu", "Tính chất quan trọng", "Yếu tố cốt lõi", "Khía cạnh đặc biệt"],
            ["Hướng tiếp cận mới", "Góc nhìn đa chiều", "Phương pháp tổng hợp", "Cách thức ứng dụng"],
            ["Nguyên tắc chỉ đạo", "Quy luật cơ bản", "Chuẩn mức đánh giá", "Tiêu chí quan trọng"],
            ["Kỹ thuật chuyên biệt", "Công cụ hỗ trợ", "Phương tiện thực hiện", "Biện pháp cải thiện"]
        ]
        
        question = question_templates[variation % len(question_templates)]
        answer_set = answer_pools[variation % len(answer_pools)]
        
        choices = []
        for i, answer in enumerate(answer_set):
            choices.append(f"{'ABCD'[i]}. {answer}")

        return {
            "id": q_id,
            "question": question,
            "choices": choices,
            "correct_answer": correct_answer,
            "difficulty": difficulty
        }

    def _generate_quiz_question(self, q_id: int, difficulty: str, lessons: List[Dict]) -> Dict:
        """Tạo một câu hỏi quiz từ các bài học"""
        import random
        lesson = random.choice(lessons) if lessons else None

        if not lesson:
            return self._generate_default_quiz_question(q_id, difficulty)

        content = lesson.get("content", "")
        title = lesson.get("title", f"Bài {q_id}")
        
        # Tạo đa dạng câu hỏi và đáp án
        question_types = [
            f"Theo {title}, điểm chính được nhấn mạnh là gì?",
            f"Nội dung {title} tập trung vào vấn đề nào?", 
            f"Ý tưởng quan trọng nhất trong {title} là gì?",
            f"Thông tin cốt lõi của {title} đề cập đến?",
            f"{title} chủ yếu giải thích về điều gì?"
        ]
        
        # Tạo đáp án ngẫu nhiên từ nội dung
        words = content.split()
        correct_answer = random.randint(0, 3)
        
        choices = []
        for i in range(4):
            if i == correct_answer and len(words) >= 8:
                # Đáp án đúng từ nội dung thực
                answer = " ".join(words[:8])
                if len(answer) > 65:
                    answer = answer[:65]
                choices.append(f"{'ABCD'[i]}. {answer}")
            else:
                # Đáp án sai đa dạng
                fake_answers = [
                    "Khái niệm lý thuyết cơ bản",
                    "Phương pháp thực hành",
                    "Ví dụ minh họa cụ thể", 
                    "Kết luận và tổng kết",
                    "Nguyên tắc áp dụng",
                    "Quy trình thực hiện",
                    "Phân tích chi tiết",
                    "Đánh giá tổng quan"
                ]
                answer = random.choice(fake_answers)
                choices.append(f"{'ABCD'[i]}. {answer}")

        return {
            "id": q_id,
            "question": random.choice(question_types),
            "choices": choices,
            "correct_answer": correct_answer,
            "difficulty": difficulty
        }

    def _generate_default_quiz_question(self, q_id: int, difficulty: str) -> Dict:
        """Tạo câu hỏi quiz mặc định"""
        correct_ans = random.randint(0, 3)
        return {
            "id": q_id,
            "question": f"Câu hỏi {q_id} về nội dung học ({difficulty})",
            "choices": [
                "A. Khái niệm cơ bản", 
                "B. Phương pháp áp dụng", 
                "C. Ví dụ thực tế", 
                "D. Kết luận chung"
            ],
            "correct_answer": correct_ans,
            "difficulty": difficulty
        }

    def _create_fallback_quiz(self, lessons: List[Dict], total_questions: int) -> Dict:
        """Tạo quiz dự phòng từ lessons với tránh trùng lặp"""
        quiz = {"easy": [], "medium": [], "hard": []}
        
        easy_count = int(total_questions * 0.4)
        medium_count = int(total_questions * 0.4)
        hard_count = total_questions - easy_count - medium_count

        question_id = 1

        # Tạo câu hỏi easy
        for i in range(easy_count):
            question = self._generate_unique_quiz_question(question_id, "easy", lessons)
            if question:
                quiz["easy"].append(question)
                question_id += 1
            
        # Tạo câu hỏi medium
        for i in range(medium_count):
            question = self._generate_unique_quiz_question(question_id, "medium", lessons)
            if question:
                quiz["medium"].append(question)
                question_id += 1
            
        # Tạo câu hỏi hard
        for i in range(hard_count):
            question = self._generate_unique_quiz_question(question_id, "hard", lessons)
            if question:
                quiz["hard"].append(question)
                question_id += 1
            
        return quiz

    # =========================
    # Lưu dữ liệu an toàn
    # =========================

    def _safe_save_data_files(self, lessons: List[Dict], quiz: Dict) -> bool:
        """Lưu dữ liệu vào file với error handling toàn diện"""
        try:
            # Tính toán thống kê
            total_lesson_questions = sum(len(lesson.get('questions', [])) for lesson in lessons)
            total_quiz_questions = sum(len(q) for q in quiz.values()) if isinstance(quiz, dict) else 0

            # Tạo backup paths
            lessons_backup = self.lessons_path + '.backup'
            quiz_backup = self.quiz_path + '.backup'

            # Backup files cũ nếu tồn tại
            if os.path.exists(self.lessons_path):
                try:
                    import shutil
                    shutil.copy2(self.lessons_path, lessons_backup)
                except Exception:
                    pass  # Ignore backup errors

            if os.path.exists(self.quiz_path):
                try:
                    import shutil
                    shutil.copy2(self.quiz_path, quiz_backup)
                except Exception:
                    pass  # Ignore backup errors

            # Chuẩn bị dữ liệu lessons
            lessons_data = {
                "metadata": {
                    "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                    "total_lessons": len(lessons),
                    "total_questions": total_lesson_questions,
                    "questions_per_lesson": self.questions_per_lesson,
                    "unique_questions": len(self.used_questions),
                    "settings": {
                        "lessons_count": self.lessons_count,
                        "questions_per_lesson": self.questions_per_lesson,
                        "quiz_questions": self.quiz_questions
                    }
                },
                "lessons": lessons
            }

            # Chuẩn bị dữ liệu quiz
            quiz_data = {
                "metadata": {
                    "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                    "total_questions": total_quiz_questions,
                    "unique_quiz_questions": len(self.used_quiz_questions),
                    "difficulty_distribution": {
                        "easy": len(quiz.get("easy", [])),
                        "medium": len(quiz.get("medium", [])),
                        "hard": len(quiz.get("hard", []))
                    }
                },
                **quiz
            }

            # Kiểm tra JSON serializable trước khi lưu
            try:
                json.dumps(lessons_data, ensure_ascii=False)
                json.dumps(quiz_data, ensure_ascii=False)
            except (TypeError, ValueError) as e:
                return False

            # Lưu lessons với atomic write
            temp_lessons_path = self.lessons_path + '.tmp'
            with open(temp_lessons_path, 'w', encoding='utf-8') as f:
                json.dump(lessons_data, f, ensure_ascii=False, indent=2)
            
            # Atomic rename
            if os.name == 'nt':  # Windows
                if os.path.exists(self.lessons_path):
                    os.remove(self.lessons_path)
            os.rename(temp_lessons_path, self.lessons_path)

            # Lưu quiz với atomic write
            temp_quiz_path = self.quiz_path + '.tmp'
            with open(temp_quiz_path, 'w', encoding='utf-8') as f:
                json.dump(quiz_data, f, ensure_ascii=False, indent=2)
            
            # Atomic rename
            if os.name == 'nt':  # Windows
                if os.path.exists(self.quiz_path):
                    os.remove(self.quiz_path)
            os.rename(temp_quiz_path, self.quiz_path)

            # Kiểm tra files đã được tạo và có dữ liệu
            if not os.path.exists(self.lessons_path) or os.path.getsize(self.lessons_path) == 0:
                raise Exception("File lessons không được tạo hoặc rỗng")

            if not os.path.exists(self.quiz_path) or os.path.getsize(self.quiz_path) == 0:
                raise Exception("File quiz không được tạo hoặc rỗng")

            # Xóa backup files nếu save thành công
            try:
                if os.path.exists(lessons_backup):
                    os.remove(lessons_backup)
                if os.path.exists(quiz_backup):
                    os.remove(quiz_backup)
            except Exception:
                pass  # Ignore cleanup errors

            return True

        except Exception as e:
            
            # Restore from backup nếu có
            try:
                if os.path.exists(lessons_backup):
                    import shutil
                    shutil.copy2(lessons_backup, self.lessons_path)
                if os.path.exists(quiz_backup):
                    import shutil
                    shutil.copy2(quiz_backup, self.quiz_path)
            except Exception:
                pass

            return False

    # =========================
    # Utility methods
    # =========================

    def stop_processing(self):
        """Dừng quá trình xử lý"""
        self.should_stop = True

    def get_progress_info(self) -> Dict:
        """Lấy thông tin tiến trình"""
        return {
            "lessons_count": self.lessons_count,
            "questions_per_lesson": self.questions_per_lesson,
            "quiz_questions": self.quiz_questions,
            "timeout": self.timeout,
            "total_timeout": self.total_timeout,
            "unique_lesson_questions": len(self.used_questions),
            "unique_quiz_questions": len(self.used_quiz_questions)
        }

    def validate_config(self) -> Tuple[bool, str]:
        """Kiểm tra cấu hình hệ thống"""
        try:
            # Kiểm tra thư mục assets
            if not os.path.exists(config.ASSETS_DIR):
                try:
                    os.makedirs(config.ASSETS_DIR)
                except Exception as e:
                    return False, f"Không thể tạo thư mục {config.ASSETS_DIR}: {e}"

            # Kiểm tra quyền ghi
            test_file = os.path.join(config.ASSETS_DIR, 'test_write.tmp')
            try:
                with open(test_file, 'w', encoding='utf-8') as f:
                    f.write('test')
                os.remove(test_file)
            except Exception as e:
                return False, f"Không có quyền ghi trong {config.ASSETS_DIR}: {e}"

            # Kiểm tra chardet package
            try:
                import chardet
            except ImportError:
                return False, "Cần cài đặt chardet để detect encoding: pip install chardet"

            # Kiểm tra Gemini client
            try:
                if not hasattr(self.gemini_client, 'api_key') or not self.gemini_client.api_key:
                    return False, "Gemini API key không hợp lệ"
                if not hasattr(self.client, 'chat'):
                    return False, "Client AI không hợp lệ"
            except Exception as e:
                return False, f"Lỗi Gemini client: {e}"

            # Test Gemini connection (optional quick test)
            try:
                test_response = self.client.chat.completions.create(
                    messages=[{"role": "user", "content": "Test connection"}],
                    temperature=0.1
                )
                if not test_response or not hasattr(test_response, 'choices'):
                    return False, "Gemini API không phản hồi đúng format"
            except Exception as e:
                return False, f"Gemini API connection error: {e}"

            return True, "Cấu hình hệ thống OK - Gemini AI đã sẵn sàng với khả năng tránh trùng lặp câu hỏi"

        except Exception as e:
            return False, f"Lỗi kiểm tra cấu hình: {e}"

    def cleanup_cache(self):
        """Dọn dẹp cache và reset tracking"""
        self.content_cache.clear()
        self.used_questions.clear()
        self.used_quiz_questions.clear()

    def get_cache_info(self) -> Dict:
        """Lấy thông tin cache và tracking"""
        return {
            "content_cache_size": len(self.content_cache),
            "content_cache_keys": list(self.content_cache.keys())[:5],  # Show first 5 keys only
            "unique_lesson_questions": len(self.used_questions),
            "unique_quiz_questions": len(self.used_quiz_questions)
        }

    def test_gemini_connection(self) -> Tuple[bool, str]:
        """Test kết nối Gemini API"""
        try:
            test_prompt = "Xin chào, bạn có thể trả lời câu hỏi này không?"
            response = self.client.chat.completions.create(
                messages=[{"role": "user", "content": test_prompt}],
                temperature=0.3
            )
            
            if response and hasattr(response, 'choices') and len(response.choices) > 0:
                content = response.choices[0].message.content
                if content and len(content.strip()) > 0:
                    return True, f"Gemini AI hoạt động tốt. Response: {content[:100]}..."
                else:
                    return False, "Gemini API trả về response rỗng"
            else:
                return False, "Gemini API không trả về response hợp lệ"
                
        except Exception as e:
            return False, f"Lỗi kết nối Gemini: {e}"

    def get_ai_info(self) -> Dict:
        """Lấy thông tin AI đang sử dụng"""
        return {
            "ai_provider": "Google Gemini",
            "model": "gemini-2.0-flash",
            "api_key_prefix": self.gemini_client.api_key[:10] + "..." if self.gemini_client.api_key else "None",
            "base_url": self.gemini_client.base_url,
            "timeout": self.timeout,
            "max_retries": self.max_retries,
            "features": [
                "Auto encoding detection",
                "UTF-16 support", 
                "Question deduplication",
                "Intelligent chunking",
                "Atomic file operations"
            ]
        }

    def test_encoding_detection(self, file_path: str) -> Tuple[bool, str]:
        """Test khả năng detect encoding của file"""
        try:
            if not os.path.exists(file_path):
                return False, "File không tồn tại"
            
            detected_encoding = self._detect_encoding(file_path)
            
            # Thử đọc file với encoding đã detect
            try:
                with open(file_path, 'r', encoding=detected_encoding) as f:
                    content = f.read(100)  # Đọc 100 ký tự đầu
                    return True, f"Detected encoding: {detected_encoding}. Sample: {content[:50]}..."
            except Exception as e:
                return False, f"Detected encoding {detected_encoding} nhưng không đọc được: {e}"
                
        except Exception as e:
            return False, f"Lỗi test encoding detection: {e}"

    def get_supported_encodings(self) -> List[str]:
        """Lấy danh sách encoding được hỗ trợ"""
        return [
            'utf-8', 'utf-16', 'utf-16-le', 'utf-16-be',
            'cp1252', 'latin-1', 'gbk', 'big5',
            'windows-1252', 'iso-8859-1'
        ]

    def analyze_file_info(self, file_path: str) -> Dict:
        """Phân tích thông tin file trước khi xử lý"""
        try:
            if not os.path.exists(file_path):
                return {"error": "File không tồn tại"}
            
            file_size = os.path.getsize(file_path)
            file_ext = os.path.splitext(file_path)[1].lower()
            
            info = {
                "file_path": file_path,
                "file_size": file_size,
                "file_extension": file_ext,
                "is_supported": self.is_supported(file_path),
                "estimated_processing_time": self._estimate_processing_time(file_size),
            }
            
            # Detect encoding cho text files
            if file_ext == '.txt':
                try:
                    detected_encoding = self._detect_encoding(file_path)
                    info["detected_encoding"] = detected_encoding
                    
                    # Test đọc với encoding đã detect
                    with open(file_path, 'r', encoding=detected_encoding) as f:
                        sample = f.read(200)
                        info["sample_content"] = sample[:100] + "..." if len(sample) > 100 else sample
                        info["encoding_confidence"] = "high" if not self._is_binary_garbage(sample) else "low"
                except Exception as e:
                    info["encoding_error"] = str(e)
            
            return info
            
        except Exception as e:
            return {"error": f"Lỗi phân tích file: {e}"}

    def _estimate_processing_time(self, file_size: int) -> str:
        """Ước tính thời gian xử lý dựa trên kích thước file"""
        if file_size < 1024:  # < 1KB
            return "< 30s"
        elif file_size < 10 * 1024:  # < 10KB
            return "30s - 1m"
        elif file_size < 100 * 1024:  # < 100KB
            return "1m - 3m"
        elif file_size < 1024 * 1024:  # < 1MB
            return "3m - 5m"
        else:
            return "> 5m"

    def reset_tracking(self):
        """Reset tracking sets để bắt đầu fresh cho file mới"""
        self.used_questions.clear()
        self.used_quiz_questions.clear()

    def get_duplication_stats(self) -> Dict:
        """Lấy thống kê về việc tránh trùng lặp"""
        return {
            "total_lesson_questions_generated": len(self.used_questions),
            "total_quiz_questions_generated": len(self.used_quiz_questions),
            "deduplication_enabled": True,
            "hash_algorithm": "MD5",
            "tracking_fields": ["question_text", "answer_choices"]
        }